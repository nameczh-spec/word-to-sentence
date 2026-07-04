import os
import json
import re
from docx import Document

WORD_DIR = r'h:\IT\单词化句\语料库\word版英语真题'
JSON_DIR = r'h:\IT\单词化句\vocab-reading-generator\vocab-lib\pastpapers'


def table_to_html(table):
    html = '<table class="writing-table">\n'
    for i, row in enumerate(table.rows):
        tag = 'th' if i == 0 else 'td'
        html += '  <tr>\n'
        for cell in row.cells:
            text = cell.text.strip()
            text = re.sub(r'\n+', '<br>', text)
            html += f'    <{tag}>{text}</{tag}>\n'
        html += '  </tr>\n'
    html += '</table>'
    return html


def extract_writing_from_docx(docx_path):
    doc = Document(docx_path)
    
    writing_section_start = -1
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        if 'Section IV' in text and 'Writing' in text:
            writing_section_start = i
            break
        if 'Section Ⅳ' in text and 'Writing' in text:
            writing_section_start = i
            break
    
    if writing_section_start < 0:
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if ('Section III' in text or 'Section Ⅲ' in text) and 'Writing' in text and i > 50:
                writing_section_start = i
                break
    
    if writing_section_start < 0:
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if 'Writing' in text and i > 100 and len(text) < 50:
                writing_section_start = i
                break
    
    if writing_section_start < 0:
        return None
    
    writing_part_a_start = -1
    for i in range(writing_section_start, min(writing_section_start + 30, len(doc.paragraphs))):
        text = doc.paragraphs[i].text.strip()
        if 'Part A' in text:
            writing_part_a_start = i
            break
    
    if writing_part_a_start < 0:
        writing_part_a_start = writing_section_start + 1
    
    table_map = {}
    for tbl in doc.tables:
        prev = tbl._element.getprevious()
        table_map[prev] = tbl
    
    part_a_parts = []
    part_b_parts = []
    current_part = 'A'
    part_b_started = False
    
    for i in range(writing_part_a_start, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        text = para.text.strip()
        
        if '答案' in text and len(text) < 50:
            break
        if 'Section I' in text and i > writing_part_a_start + 20:
            break
        if 'Section Ⅰ' in text and i > writing_part_a_start + 20:
            break
        
        if not part_b_started and 'Part B' in text:
            current_part = 'B'
            part_b_started = True
            if text.strip() == 'Part B':
                continue
        
        para_elem = para._element
        if para_elem in table_map:
            table = table_map[para_elem]
            if current_part == 'B' and part_b_started:
                part_b_parts.append(table_to_html(table))
            elif current_part == 'A':
                part_a_parts.append(table_to_html(table))
        
        if text:
            if current_part == 'A':
                part_a_parts.append(text)
            else:
                part_b_parts.append(text)
    
    part_a_text = '\n'.join(part_a_parts).strip()
    part_b_text = '\n'.join(part_b_parts).strip()
    
    def clean_page_numbers(text):
        lines = text.split('\n')
        cleaned = []
        for line in lines:
            stripped = line.strip()
            if re.match(r'^\d+$', stripped) and len(stripped) <= 3:
                continue
            if re.match(r'^[.·\s]*\d{1,3}[.·\s]*$', stripped):
                continue
            if re.match(r'^英语试题', stripped):
                continue
            if re.match(r'^第\s*\d+\s*页', stripped):
                continue
            if re.match(r'^\d{4}$', stripped):
                continue
            if re.match(r'^年英语[一二]真题', stripped):
                continue
            if re.match(r'^共\s*\d+\s*页', stripped):
                continue
            cleaned.append(line)
        result = '\n'.join(cleaned).strip()
        result = re.sub(r'\n{3,}', '\n\n', result)
        return result
    
    part_a_text = clean_page_numbers(part_a_text)
    part_b_text = clean_page_numbers(part_b_text)
    
    if 'Read the following four texts' in part_a_text or 'Read the following text' in part_a_text:
        part_a_text = ''
    
    if 'Read the following four texts' in part_b_text or 'Read the following text' in part_b_text:
        part_b_text = part_b_text.split('Read the following')[0].strip()
    
    return {
        'partA': part_a_text,
        'partB': part_b_text
    }


def main():
    word_files = sorted([f for f in os.listdir(WORD_DIR) if f.endswith('.docx')])
    
    updated = []
    skipped = []
    errors = []
    
    for fname in word_files:
        paper_id = fname.replace('.docx', '')
        json_path = os.path.join(JSON_DIR, f'{paper_id}.json')
        
        if not os.path.exists(json_path):
            skipped.append(f'{paper_id}: JSON不存在')
            continue
        
        docx_path = os.path.join(WORD_DIR, fname)
        
        try:
            content = extract_writing_from_docx(docx_path)
            if content is None:
                skipped.append(f'{paper_id}: 未找到作文部分')
                continue
            
            with open(json_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            writing_section = None
            for sec in data.get('sections', []):
                if sec.get('type') == 'writing':
                    writing_section = sec
                    break
            
            if writing_section is None:
                skipped.append(f'{paper_id}: JSON无writing section')
                continue
            
            old_content = writing_section.get('content', '')
            
            if content['partB']:
                new_full_content = content['partA'] + '\n\nPart B\n' + content['partB']
            else:
                new_full_content = content['partA']
            new_full_content = new_full_content.strip()
            
            has_read_error = 'Read the following four texts' in old_content or 'Read the following text' in old_content
            has_partb_missing = 'Part B' not in old_content and content['partB']
            
            if has_read_error or has_partb_missing or new_full_content.strip() != old_content.strip():
                writing_section['content'] = new_full_content
                
                with open(json_path, 'w', encoding='utf-8') as f:
                    json.dump(data, f, ensure_ascii=False, indent=2)
                
                if has_read_error:
                    updated.append(f'{paper_id}: 已修复阅读题混入问题')
                elif has_partb_missing:
                    updated.append(f'{paper_id}: 已添加Part B标记')
                else:
                    updated.append(f'{paper_id}: 已更新')
            else:
                skipped.append(f'{paper_id}: 内容正确，无需修改')
        
        except Exception as e:
            errors.append(f'{paper_id}: 错误 - {e}')
            import traceback
            traceback.print_exc()
    
    print('=== 已更新 ===')
    for s in updated:
        print(s)
    
    print(f'\n总计：更新 {len(updated)} 个，跳过 {len(skipped)} 个，错误 {len(errors)} 个')
    
    if errors:
        print('\n=== 错误 ===')
        for s in errors:
            print(s)


if __name__ == '__main__':
    main()
