import os
import json
from docx import Document
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT

WORD_DIR = r'h:\IT\单词化句\语料库\word版英语真题'
JSON_DIR = r'h:\IT\单词化句\vocab-reading-generator\vocab-lib\pastpapers'
OUTPUT_DIR = r'h:\IT\单词化句\.dev-tools\writing-extracted'

os.makedirs(OUTPUT_DIR, exist_ok=True)


def table_to_html(table):
    html = '<table class="writing-table">\n'
    for i, row in enumerate(table.rows):
        tag = 'th' if i == 0 else 'td'
        html += '  <tr>\n'
        for cell in row.cells:
            text = cell.text.strip()
            html += f'    <{tag}>{text}</{tag}>\n'
        html += '  </tr>\n'
    html += '</table>'
    return html


def is_writing_start(para_text, para_index):
    text = para_text.strip()
    if para_index < 50:
        return False
    if 'Part A' in text and ('Directions' in text or 'direction' in text.lower()):
        return True
    if text.startswith('Part A') and len(text) < 50:
        return True
    return False


def extract_images_from_doc(doc, output_dir, paper_id):
    image_dir = os.path.join(output_dir, 'images')
    os.makedirs(image_dir, exist_ok=True)
    
    images = []
    img_idx = 0
    
    for rel in doc.part.rels.values():
        if 'image' in rel.reltype:
            image_part = rel.target_part
            img_bytes = image_part.blob
            ext = image_part.content_type.split('/')[-1]
            if ext == 'jpeg':
                ext = 'jpg'
            
            img_idx += 1
            img_name = f'{paper_id}-img{img_idx}.{ext}'
            img_path = os.path.join(image_dir, img_name)
            
            with open(img_path, 'wb') as f:
                f.write(img_bytes)
            
            images.append({
                'name': img_name,
                'path': img_path,
                'size': len(img_bytes),
                'rid': rel.rId
            })
    
    return images


def get_paragraph_image_rids(paragraph):
    rids = []
    elem = paragraph._element
    xml_str = elem.xml
    
    import re
    embed_matches = re.findall(r'r:embed="([^"]+)"', xml_str)
    rids.extend(embed_matches)
    
    id_matches = re.findall(r'r:id="([^"]+)"', xml_str)
    rids.extend(id_matches)
    
    return list(set(rids))


def extract_writing_from_docx(docx_path, paper_id):
    doc = Document(docx_path)
    
    writing_start = -1
    for i, para in enumerate(doc.paragraphs):
        if is_writing_start(para.text, i):
            writing_start = i
            break
    
    if writing_start < 0:
        return None
    
    images = extract_images_from_doc(doc, OUTPUT_DIR, paper_id)
    rid_to_image = {}
    for img in images:
        rid_to_image[img['rid']] = img
    
    content_parts = []
    table_idx = 0
    all_tables = doc.tables
    
    for i in range(writing_start, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        text = para.text.strip()
        
        img_rids = get_paragraph_image_rids(para)
        for rid in img_rids:
            if rid in rid_to_image:
                img = rid_to_image[rid]
                content_parts.append({
                    'type': 'image',
                    'src': img['name'],
                    'size': img['size']
                })
        
        if text:
            content_parts.append({
                'type': 'text',
                'text': text
            })
        
        next_para_idx = i + 1
        if next_para_idx < len(doc.paragraphs):
            next_para_element = doc.paragraphs[next_para_idx]._element
            for tbl in all_tables:
                tbl_element = tbl._element
                if next_para_element.getnext() is not None and next_para_element.getnext() == tbl_element:
                    pass
    
    table_xml_map = {}
    for tbl in all_tables:
        tbl_xml = tbl._element
        prev = tbl_xml.getprevious()
        table_xml_map[prev] = tbl
    
    results = []
    for i in range(writing_start, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        text = para.text.strip()
        
        img_rids = get_paragraph_image_rids(para)
        for rid in img_rids:
            if rid in rid_to_image:
                img = rid_to_image[rid]
                results.append(f'[IMAGE] {img["name"]} ({img["size"]} bytes)')
        
        if text:
            results.append(text)
        
        para_element = para._element
        if para_element in table_xml_map:
            table = table_xml_map[para_element]
            results.append('[TABLE]')
            results.append(table_to_html(table))
    
    part_a_text = []
    part_b_text = []
    current = 'A'
    
    for item in results:
        if isinstance(item, str) and 'Part B' in item and current == 'A':
            current = 'B'
        
        if current == 'A':
            part_a_text.append(item)
        else:
            part_b_text.append(item)
    
    return {
        'paper_id': paper_id,
        'partA': part_a_text,
        'partB': part_b_text,
        'images': images,
        'full_text': '\n'.join([str(x) for x in results])
    }


def main():
    word_files = sorted([f for f in os.listdir(WORD_DIR) if f.endswith('.docx')])
    
    summary = []
    
    for fname in word_files:
        paper_id = fname.replace('.docx', '')
        docx_path = os.path.join(WORD_DIR, fname)
        
        try:
            result = extract_writing_from_docx(docx_path, paper_id)
            
            if result is None:
                summary.append(f'{paper_id}: 未找到作文部分')
                continue
            
            json_path = os.path.join(JSON_DIR, f'{paper_id}.json')
            has_json = os.path.exists(json_path)
            
            has_table_a = any('[TABLE]' in str(x) for x in result['partA'])
            has_table_b = any('[TABLE]' in str(x) for x in result['partB'])
            has_image_a = any('[IMAGE]' in str(x) for x in result['partA'])
            has_image_b = any('[IMAGE]' in str(x) for x in result['partB'])
            
            img_count = len(result['images'])
            
            info = f'{paper_id}: JSON={has_json}, 图片={img_count}'
            if has_table_a:
                info += ', PartA含表格'
            if has_table_b:
                info += ', PartB含表格'
            if has_image_a:
                info += ', PartA含图片'
            if has_image_b:
                info += ', PartB含图片'
            
            summary.append(info)
            
            out_path = os.path.join(OUTPUT_DIR, f'{paper_id}-writing.txt')
            with open(out_path, 'w', encoding='utf-8') as f:
                f.write(f'=== {paper_id} 作文内容 ===\n\n')
                f.write('--- Part A ---\n')
                for item in result['partA']:
                    f.write(str(item) + '\n')
                f.write('\n--- Part B ---\n')
                for item in result['partB']:
                    f.write(str(item) + '\n')
            
        except Exception as e:
            summary.append(f'{paper_id}: 错误 - {e}')
            import traceback
            traceback.print_exc()
    
    print('=== 提取汇总 ===')
    for s in summary:
        print(s)
    
    summary_path = os.path.join(OUTPUT_DIR, 'SUMMARY.txt')
    with open(summary_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(summary))
    
    print(f'\n详细文件保存在: {OUTPUT_DIR}')


if __name__ == '__main__':
    main()
