import os
import json
import re
from docx import Document
from docx.oxml.ns import qn

WORD_DIR = r'h:\IT\单词化句\语料库\word版英语真题'
JSON_DIR = r'h:\IT\单词化句\vocab-reading-generator\vocab-lib\pastpapers'
IMG_DIR = r'h:\IT\单词化句\vocab-reading-generator\images\writing'


def table_to_html(table):
    html = '<table class="writing-table">\n'
    for i, row in enumerate(table.rows):
        tag = 'th' if i == 0 else 'td'
        html += '  <tr>\n'
        for cell in row.cells:
            text = cell.text.strip()
            text = re.sub(r'\n+', '<br>', text)
            html += f'    <{tag}>{text}</{tag}>\n'
        html += '  </tr>\n'
    html += '</table>'
    return html


def get_para_image_count(paragraph):
    elem = paragraph._element
    xml_str = elem.xml
    return len(re.findall(r'r:embed="[^"]+"', xml_str)) + len(re.findall(r'r:id="[^"]+"', xml_str))


def extract_writing_content(docx_path):
    doc = Document(docx_path)
    
    writing_start = -1
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if i < 30:
            continue
        if 'Part A' in text and 'Directions' in text and i > 50:
            writing_start = i
            break
        if text.startswith('Part A') and i > 100:
            writing_start = i
            break
    
    if writing_start < 0:
        return None
    
    table_map = {}
    for tbl in doc.tables:
        prev = tbl._element.getprevious()
        table_map[prev] = tbl
    
    part_a_lines = []
    part_b_lines = []
    current_part = 'A'
    part_b_started = False
    
    for i in range(writing_start, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        text = para.text.strip()
        
        if not part_b_started and 'Part B' in text:
            current_part = 'B'
            part_b_started = True
        
        if part_b_started and ('Section III' in text or 'Section Ⅲ' in text or '答案' in text or '参考答案' in text):
            break
        
        para_elem = para._element
        if para_elem in table_map:
            table = table_map[para_elem]
            html_table = table_to_html(table)
            if current_part == 'A':
                part_a_lines.append(html_table)
            else:
                part_b_lines.append(html_table)
        
        if text:
            if current_part == 'A':
                part_a_lines.append(text)
            else:
                part_b_lines.append(text)
    
    return {
        'partA': '\n'.join(part_a_lines),
        'partB': '\n'.join(part_b_lines)
    }


def main():
    word_files = sorted([f for f in os.listdir(WORD_DIR) if f.endswith('.docx')])
    
    updated = []
    skipped = []
    errors = []
    
    for fname in word_files:
        paper_id = fname.replace('.docx', '')
        json_path = os.path.join(JSON_DIR, f'{paper_id}.json')
        
        if not os.path.exists(json_path):
            skipped.append(f'{paper_id}: JSON不存在')
            continue
        
        docx_path = os.path.join(WORD_DIR, fname)
        
        try:
            content = extract_writing_content(docx_path)
            if content is None:
                skipped.append(f'{paper_id}: 未找到作文部分')
                continue
            
            with open(json_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            writing_section = None
            for sec in data.get('sections', []):
                if sec.get('type') == 'writing':
                    writing_section = sec
                    break
            
            if writing_section is None:
                skipped.append(f'{paper_id}: JSON无writing section')
                continue
            
            old_content = writing_section.get('content', '')
            has_table_old = '<table' in old_content
            has_table_new = '<table' in content['partB']
            
            if has_table_new and not has_table_old:
                new_content = f"{content['partA']}\n\n{content['partB']}"
                writing_section['content'] = new_content
                
                with open(json_path, 'w', encoding='utf-8') as f:
                    json.dump(data, f, ensure_ascii=False, indent=2)
                
                updated.append(f'{paper_id}: 已添加表格（PartB）')
            else:
                skipped.append(f'{paper_id}: 无需修改（旧版已有表格={has_table_old}，新版有表格={has_table_new}）')
        
        except Exception as e:
            errors.append(f'{paper_id}: 错误 - {e}')
            import traceback
            traceback.print_exc()
    
    print('=== 已更新 ===')
    for s in updated:
        print(s)
    
    print('\n=== 跳过 ===')
    for s in skipped:
        print(s)
    
    print('\n=== 错误 ===')
    for s in errors:
        print(s)
    
    print(f'\n总计：更新 {len(updated)} 个，跳过 {len(skipped)} 个，错误 {len(errors)} 个')


if __name__ == '__main__':
    main()
