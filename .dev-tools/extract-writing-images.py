import os
import re
from docx import Document
from docx.oxml.ns import qn

WORD_DIR = r'h:\IT\单词化句\语料库\word版英语真题'
IMG_DIR = r'h:\IT\单词化句\vocab-reading-generator\images\writing'


def get_paragraph_image_rids(paragraph):
    rids = []
    elem = paragraph._element
    xml_str = elem.xml
    embed_matches = re.findall(r'r:embed="([^"]+)"', xml_str)
    rids.extend(embed_matches)
    id_matches = re.findall(r'r:id="([^"]+)"', xml_str)
    rids.extend(id_matches)
    return list(set(rids))


def extract_partb_image(docx_path, paper_id):
    doc = Document(docx_path)
    
    part_b_start = -1
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if i < 30:
            continue
        if 'Part B' in text and ('Directions' in text or len(text) < 20):
            part_b_start = i
            break
    
    if part_b_start < 0:
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if 'Part B' in text and i > 50:
                part_b_start = i
                break
    
    if part_b_start < 0:
        return None
    
    rid_to_image = {}
    for rel in doc.part.rels.values():
        if 'image' in rel.reltype:
            image_part = rel.target_part
            img_bytes = image_part.blob
            ext = image_part.content_type.split('/')[-1]
            if ext == 'jpeg':
                ext = 'jpg'
            rid_to_image[rel.rId] = {
                'bytes': img_bytes,
                'ext': ext,
                'size': len(img_bytes)
            }
    
    part_b_images = []
    found_part_b = False
    
    for i in range(part_b_start, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        text = para.text.strip()
        
        if 'Section III' in text or 'Section Ⅲ' in text:
            break
        if ('答案' in text or '参考答案' in text) and len(part_b_images) > 0:
            break
        
        img_rids = get_paragraph_image_rids(para)
        for rid in img_rids:
            if rid in rid_to_image and rid not in [x['rid'] for x in part_b_images]:
                img = rid_to_image[rid]
                img['rid'] = rid
                part_b_images.append(img)
    
    if len(part_b_images) == 0 and len(rid_to_image) > 0:
        max_size = 0
        max_img = None
        for rid, img in rid_to_image.items():
            if img['size'] > max_size:
                max_size = img['size']
                max_img = img
        if max_img and max_img['size'] > 5000:
            max_img['rid'] = 'fallback'
            part_b_images.append(max_img)
    
    part_b_images = [img for img in part_b_images if img['size'] > 1000]
    
    return part_b_images


def main():
    word_files = sorted([f for f in os.listdir(WORD_DIR) if f.endswith('.docx')])
    
    extracted = []
    skipped = []
    errors = []
    
    for fname in word_files:
        paper_id = fname.replace('.docx', '')
        docx_path = os.path.join(WORD_DIR, fname)
        
        try:
            images = extract_partb_image(docx_path, paper_id)
            
            if images is None or len(images) == 0:
                skipped.append(f'{paper_id}: PartB无图片')
                continue
            
            img = images[0]
            ext = img['ext']
            
            target_name = f'{paper_id}-partb.{ext}'
            target_path = os.path.join(IMG_DIR, target_name)
            
            old_files = [f for f in os.listdir(IMG_DIR) if f.startswith(f'{paper_id}-') and 'partb' in f.lower()]
            
            with open(target_path, 'wb') as f:
                f.write(img['bytes'])
            
            info = f'{paper_id}: 提取图片 -> {target_name} ({img["size"]} bytes)'
            if old_files:
                info += f' [原有: {", ".join(old_files)}]'
            extracted.append(info)
        
        except Exception as e:
            errors.append(f'{paper_id}: 错误 - {e}')
            import traceback
            traceback.print_exc()
    
    print('=== 已提取图片 ===')
    for s in extracted:
        print(s)
    
    print(f'\n总计：提取 {len(extracted)} 个，跳过 {len(skipped)} 个，错误 {len(errors)} 个')
    
    if errors:
        print('\n=== 错误 ===')
        for s in errors:
            print(s)


if __name__ == '__main__':
    main()
