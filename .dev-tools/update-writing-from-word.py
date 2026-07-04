import os
import json
import re
from docx import Document
from docx.oxml.ns import qn

WORD_DIR = r'h:\IT\单词化句\语料库\word版英语真题'
JSON_DIR = r'h:\IT\单词化句\vocab-reading-generator\vocab-lib\pastpapers'


def table_to_html(table):
    html = '<table class="writing-table">\n'
    for i, row in enumerate(table.rows):
        tag = 'th' if i == 0 else 'td'
        html += '  <tr>\n'
        for cell in row.cells:
            text = cell.text.strip()
            text = re.sub(r'\n+', '<br>', text)
            html += f'    <{tag}>{text}</{tag}>\n'
        html += '  </tr>\n'
    html += '</table>'
    return html


def extract_writing_from_docx(docx_path):
    doc = Document(docx_path)
    
    writing_start = -1
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if i < 30:
            continue
        if 'Part A' in text and 'Directions' in text and i > 50:
            writing_start = i
            break
        if text.startswith('Part A') and i > 100:
            writing_start = i
            break
    
    if writing_start < 0:
        return None
    
    table_map = {}
    for tbl in doc.tables:
        prev = tbl._element.getprevious()
        table_map[prev] = tbl
    
    part_a_parts = []
    part_b_parts = []
    current_part = 'A'
    part_b_started = False
    
    for i in range(writing_start, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        text = para.text.strip()
        
        if not part_b_started and 'Part B' in text:
            current_part = 'B'
            part_b_started = True
        
        if part_b_started and ('Section III' in text or 'Section Ⅲ' in text):
            break
        if part_b_started and ('答案' in text or '参考答案' in text or 'Section II' in text):
            if len(part_b_parts) > 3:
                break
        
        para_elem = para._element
        if para_elem in table_map:
            table = table_map[para_elem]
            if current_part == 'B' and part_b_started:
                part_b_parts.append(table_to_html(table))
            elif current_part == 'A':
                part_a_parts.append(table_to_html(table))
        
        if text:
            if current_part == 'A':
                part_a_parts.append(text)
            else:
                part_b_parts.append(text)
    
    return {
        'partA': '\n'.join(part_a_parts),
        'partB': '\n'.join(part_b_parts)
    }


def get_paper_year_type(paper_id):
    m = re.match(r'(\d{4})-(1|2)', paper_id)
    if m:
        return m.group(1), '英语一' if m.group(2) == '1' else '英语二'
    return '', ''


def main():
    word_files = sorted([f for f in os.listdir(WORD_DIR) if f.endswith('.docx')])
    
    updated = []
    skipped = []
    errors = []
    
    for fname in word_files:
        paper_id = fname.replace('.docx', '')
        json_path = os.path.join(JSON_DIR, f'{paper_id}.json')
        
        if not os.path.exists(json_path):
            skipped.append(f'{paper_id}: JSON不存在')
            continue
        
        docx_path = os.path.join(WORD_DIR, fname)
        
        try:
            content = extract_writing_from_docx(docx_path)
            if content is None:
                skipped.append(f'{paper_id}: 未找到作文部分')
                continue
            
            with open(json_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            writing_section = None
            for sec in data.get('sections', []):
                if sec.get('type') == 'writing':
                    writing_section = sec
                    break
            
            if writing_section is None:
                skipped.append(f'{paper_id}: JSON无writing section')
                continue
            
            old_content = writing_section.get('content', '')
            has_table_old = '<table' in old_content
            has_table_new = '<table' in content['partB']
            
            new_full_content = content['partA'] + '\n\n' + content['partB']
            
            if new_full_content.strip() != old_content.strip():
                writing_section['content'] = new_full_content
                
                with open(json_path, 'w', encoding='utf-8') as f:
                    json.dump(data, f, ensure_ascii=False, indent=2)
                
                change_info = ''
                if has_table_new and not has_table_old:
                    change_info = '（新增表格）'
                elif has_table_new and has_table_old:
                    change_info = '（更新内容）'
                updated.append(f'{paper_id}: 已更新{change_info}')
            else:
                skipped.append(f'{paper_id}: 内容一致，无需修改')
        
        except Exception as e:
            errors.append(f'{paper_id}: 错误 - {e}')
            import traceback
            traceback.print_exc()
    
    print('=== 已更新 ===')
    for s in updated:
        print(s)
    
    print(f'\n总计：更新 {len(updated)} 个，跳过 {len(skipped)} 个，错误 {len(errors)} 个')
    
    if errors:
        print('\n=== 错误 ===')
        for s in errors:
            print(s)


if __name__ == '__main__':
    main()
